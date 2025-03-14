from docx import Document
import re
from typing import Dict, Any
import io

class ResumeParser:
    def __init__(self):
        self.sections = {
            'skills': r'SKILLS|TECHNICAL SKILLS|EXPERTISE',
            'experience': r'EXPERIENCE|WORK EXPERIENCE|EMPLOYMENT',
            'education': r'EDUCATION|ACADEMIC|QUALIFICATIONS'
        }

    def parse(self, file) -> Dict[str, Any]:
        """Parse resume file and extract relevant information."""
        try:
            # Handle both file-like objects and bytes
            if hasattr(file, 'read'):
                doc = Document(file)
            else:
                doc = Document(io.BytesIO(file))
                
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            parsed_data = {
                'skills': self._extract_skills(text),
                'experience': self._extract_experience(text),
                'education': self._extract_education(text),
                'raw_text': text
            }
            
            return parsed_data
        except Exception as e:
            raise Exception(f"Error parsing resume: {str(e)}")

    def _extract_skills(self, text: str) -> list:
        """Extract skills from resume text."""
        skills_section = self._get_section(text, self.sections['skills'])
        if skills_section:
            # Split by common delimiters and clean up
            skills = re.split(r'[,•|\n]', skills_section)
            return [skill.strip() for skill in skills if skill.strip()]
        return []

    def _extract_experience(self, text: str) -> list:
        """Extract work experience from resume text."""
        experience_section = self._get_section(text, self.sections['experience'])
        if experience_section:
            # Split by date patterns to separate different positions
            experiences = re.split(r'\d{4}\s*[-–]\s*(\d{4}|present)', experience_section)
            return [exp.strip() for exp in experiences if exp and exp.strip()]
        return []

    def _extract_education(self, text: str) -> list:
        """Extract education information from resume text."""
        education_section = self._get_section(text, self.sections['education'])
        if education_section:
            # Split by common education degree indicators
            education = re.split(r'(?i)(bachelor|master|phd|diploma)', education_section)
            return [edu.strip() for edu in education if edu.strip()]
        return []

    def _get_section(self, text: str, section_pattern: str) -> str:
        """Extract a specific section from the resume text."""
        pattern = re.compile(f"(?i){section_pattern}.*?(?=(SKILLS|EXPERIENCE|EDUCATION|$))", re.DOTALL)
        match = pattern.search(text)
        return match.group(0) if match else ""